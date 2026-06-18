import pandas as pd
import os
import re
import json
from docx import Document
from docx.shared import RGBColor, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from config.paths import PAPER_FOILS

class ReportEngine:
    def __init__(self, db_manager):
        self.db = db_manager
        self.custom_foils = {}
        
        # Wczytanie słownika z niestandardowymi foliami JSON
        try:
            with open(PAPER_FOILS, 'r', encoding='utf-8') as f:
                raw_foils = json.load(f)
                # Zabezpieczenie: odcinamy ewentualne zera i spacje z kluczy w samym JSON-ie
                self.custom_foils = {str(k).strip().lstrip('0'): v for k, v in raw_foils.items()}
        except Exception as e:
            print(f"Brak lub błąd pliku paper_foils.json: {e}")

    # --- POMOCNICZA METODA DO WSKAZANIA ODPowiedniego STEROWNIKA ODBC ---
    def load_excel_data(self, file_path: str) -> pd.DataFrame:
        """Wczytuje plik z Hydry (.csv lub .xlsx), automatycznie dopasowując ustawienia."""
        _, ext = os.path.splitext(file_path.lower())
        df = pd.DataFrame()
        
        if ext in ['.xlsx', '.xls']:
            try:
                df = pd.read_excel(file_path, header=1)
            except Exception as e:
                print(f"Error reading Excel: {e}")
                return pd.DataFrame()
        elif ext == '.csv':
            configs = [
                {'encoding': 'cp1250', 'sep': ';'},
                {'encoding': 'cp1250', 'sep': ','},
                {'encoding': 'utf-8', 'sep': ';'},
                {'encoding': 'utf-8', 'sep': ','},
            ]
            for config in configs:
                try:
                    temp_df = pd.read_csv(file_path, header=1, encoding=config['encoding'], sep=config['sep'], on_bad_lines='skip')
                    if 'Artykuł' in temp_df.columns:
                        df = temp_df
                        break
                except Exception:
                    continue
        
        if not df.empty and 'Artykuł' in df.columns:
            df = df.dropna(subset=['Artykuł'])
        return df

    # --- POMOCNICZA METODA DO OBLICZANIA POZOSTAŁYCH ILOŚCI (PCS) ---
    def get_bom_details(self, matnr_list: list) -> pd.DataFrame:
        """Pobiera BOM z Kronosa. Rozszerzono filtry, aby łapać wszystkie typy folii."""
        if not matnr_list:
            return pd.DataFrame()

        matnr_str = "', '".join([str(m) for m in matnr_list])
        
        # Pobieramy szerszy zakres, aby uwzględnić różne typy folii ochronnych (POSNR 0050, 0060)
        sql = f"""
        SELECT MATNR, KOLOR, IDNRK, POSNR
        FROM tblHANAIndeksBomLinia
        WHERE MATNR IN ('{matnr_str}')
          AND (
              IDNRK LIKE 'F%' 
              OR POSNR IN ('0050', '0060', '0090')
              OR (POSNR IN ('0020', '0030', '0070') AND IDNRK LIKE '0000%')
          )
        ORDER BY MATNR, POSNR
        """
        
        try:
            with self.db.raporty_engine.connect() as connection:
                df = pd.read_sql(sql, connection)
            return df
        except Exception as e:
            print(f"SQL Error (Kronos): {e}")
            return pd.DataFrame()

    # --- POMOCNICZA METODA DO SPRAWDZENIA, CZY MASZYNA JEST OBUSTRONNA ---
    def _extract_width_and_type(self, idnrk: str):
        """Dynamicznie wyciąga typ i szerokość z IDNRK. Folie numeryczne zrzuca na koniec."""
        idnrk = str(idnrk).strip()
        
        # Odcięcie wiodących zer
        clean_id = idnrk.lstrip('0')
        
        # Jeśli po odcięciu zer zostały same cyfry (nasza niestandardowa folia)
        if clean_id.isdigit():
            return 'Z_SPECIAL', 9999 
            
        parts = idnrk.split('.')
        foil_prefix = parts[0]
        try:
            width = int(parts[-1])
        except (ValueError, IndexError):
            width = 0
            
        return foil_prefix, width

    # --- GŁÓWNA METODA GENERUJĄCA DOKUMENT WORDA ---
    def generate_word_report(self, report_data: dict, machine_name: str, output_path: str):
        doc = Document()
        
        section = doc.sections[0]
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(1.0)
        
        snap_date = report_data.get("snapshot_date", "")
        shift_info = report_data.get("shift_info", "")
        
        if shift_info:
            shift_info = re.sub(r'\(zmiana\s+(\d+)\)', r'(\1)', shift_info)
            
        self._add_page_numbering(doc, machine_name, shift_info)
        
        header_text = f"{machine_name}"
        if shift_info:
            header_text += f" - {shift_info}"
        if snap_date:
            header_text += f"                {snap_date}"
             
        def draw_main_header(document):
            h = document.add_heading('', 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h.paragraph_format.space_after = Pt(0) 
            
            run_machine = h.add_run(machine_name)
            run_machine.font.size = Pt(26)
            run_machine.bold = True
            run_machine.underline = True
            run_machine.font.color.rgb = RGBColor(0, 0, 0)
            
            rest_text = ""
            if shift_info:
                rest_text += f" - {shift_info}"
            if snap_date:
                rest_text += f"                {snap_date}"
                
            if rest_text:
                run_rest = h.add_run(rest_text)
                run_rest.font.size = Pt(20)
                run_rest.bold = True
                run_rest.underline = False
                run_rest.font.color.rgb = RGBColor(0, 0, 0)

        draw_main_header(doc)
        
        data_dict = report_data.get("data", report_data)
        
        combined = data_dict.get('combined_side', [])
        sequence = data_dict.get('production_sequence', [])
        
        if combined:
            self._fill_decor_section(doc, combined, side_title='ZEWNĘTRZNA, WEWNĘTRZNA (KOMBAJN)')
            next_num = 1
        elif sequence:
            current_side = sequence[0].get('side_desc', 'Zewn.')
            chunk = []
            
            side_to_title = {
                'Zewn.': 'ZEWNĘTRZNA',
                'Wewn.': 'WEWNĘTRZNA',
                'Górna': 'STRONA GÓRNA'
            }
            
            for item in sequence:
                side_desc = item.get('side_desc', '')
                if side_desc == current_side:
                    chunk.append(item)
                else:
                    title = side_to_title.get(current_side, 'STRONA NIEZNANA')
                    self._fill_decor_section(doc, chunk, side_title=title)
                    
                    current_side = side_desc
                    chunk = [item]
                    
            if chunk:
                title = side_to_title.get(current_side, 'STRONA NIEZNANA')
                self._fill_decor_section(doc, chunk, side_title=title)
                
            next_num = 1
        else:
            doc.add_heading('BRAK DANYCH', level=1)
            next_num = 1

        # --- Sekcja FOLIA OCHRONNA ---
        protective = data_dict.get('protective', {})
        
        if protective:
            doc.add_page_break()
            draw_main_header(doc)
            
            doc.add_heading(f'{next_num}. Folia ochronna (SUMA ZBIORCZA)', level=1)
            
            # Agregujemy dane z powrotem do sumy globalnej, ale zbieramy przypisane geometrie
            aggregated_prot = {}
            is_nested = any(isinstance(v, dict) for v in protective.values())

            if is_nested:
                for base_geom, foils in protective.items():
                    for symbol, meters in foils.items():
                        if symbol not in aggregated_prot:
                            aggregated_prot[symbol] = {'meters': 0.0, 'geometries': set()}
                        aggregated_prot[symbol]['meters'] += meters
                        aggregated_prot[symbol]['geometries'].add(str(base_geom))
            else:
                for symbol, meters in protective.items():
                    aggregated_prot[symbol] = {'meters': meters, 'geometries': set()}

            # Tworzymy tabelę z 4 kolumnami: Indeks | Długość | Geometrie | Uwagi
            prot_table = doc.add_table(rows=0, cols=4)
            prot_table.style = 'Table Grid'
            
            # Poszerzamy tabelę o kolumnę na wpisanie numerów geometrii (Razem 17 cm)
            prot_widths = (Cm(3.0), Cm(2.5), Cm(6.5), Cm(5.0)) 
            def style_prot_row(row):
                for i, cell in enumerate(row.cells):
                    cell.width = prot_widths[i]
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(5)
                        p.paragraph_format.space_after = Pt(5)

            for symbol in sorted(aggregated_prot.keys()):
                data = aggregated_prot[symbol]
                row = prot_table.add_row()
                
                # 1. Indeks folii
                run_sym = row.cells[0].paragraphs[0].add_run(symbol)
                run_sym.bold = True
                
                # 2. Suma metrów (operatorzy nie muszą już liczyć!)
                val_str = f"{data['meters']:.1f}".replace('.', ',')
                run_m = row.cells[1].paragraphs[0].add_run(val_str)
                run_m.bold = True
                
                # 3. Lista geometrii (operatorzy wiedzą, na który wózek to idzie)
                geom_str = ", ".join(sorted(data['geometries']))
                row.cells[2].text = geom_str
                
                # 4. Puste miejsce na uwagi
                row.cells[3].text = '' 
                
                style_prot_row(row)
                
            next_num += 1

        # --- Sekcja FOLIA DEKORACYJNA ---
        decor_summary = {}
        for side in ['production_sequence', 'combined_side']:
            for item in data_dict.get(side, []):
                idnrk = item['idnrk']
                decor_summary[idnrk] = decor_summary.get(idnrk, 0.0) + item['meters']
                
        if decor_summary:
            new_section = doc.add_section()
            new_section.top_margin = Cm(0.5)
            new_section.bottom_margin = Cm(0.5)
            new_section.left_margin = Cm(1.0)
            new_section.right_margin = Cm(0.5)
            
            draw_main_header(doc)
            
            doc.add_heading(f'{next_num}. Folia dekoracyjna (SUMA ZBIORCZA)', level=1)
            
            keys = sorted(
                decor_summary.keys(), 
                key=lambda k: (self._extract_width_and_type(k)[1], self._extract_width_and_type(k)[0])
            )
            
            # --- ZMIANA: Sztywne wypełnianie lewo -> dół -> prawo ---
            MAX_ROWS_PER_PAGE = 27 # Sztywny limit wierszy na stronę
            CHUNK_SIZE = MAX_ROWS_PER_PAGE * 2 # Maksymalnie 60 folii na jednej kartce
            
            for chunk_start in range(0, len(keys), CHUNK_SIZE):
                if chunk_start > 0:
                    doc.add_page_break()
                    draw_main_header(doc)
                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(5) 
                    spacer.paragraph_format.space_after = Pt(5)
                
                chunk_keys = keys[chunk_start : chunk_start + CHUNK_SIZE]
                
                # Obliczamy ile faktycznie wierszy trzeba narysować (max 30)
                num_rows = min(MAX_ROWS_PER_PAGE, len(chunk_keys))
                
                decor_table = doc.add_table(rows=0, cols=6)
                decor_table.style = 'Table Grid'
                
                decor_widths = (Cm(2.5), Cm(2.5), Cm(5.0), Cm(2.5), Cm(2.5), Cm(5.0))

                def style_decor_row(row):
                    for i, cell in enumerate(row.cells):
                        cell.width = decor_widths[i]
                        for p in cell.paragraphs:
                            p.paragraph_format.space_before = Pt(5)
                            p.paragraph_format.space_after = Pt(5)

                # Rysujemy tabelę z góry na dół
                # Rysujemy tabelę z góry na dół
                for i in range(num_rows):
                    row = decor_table.add_row()

                    # Zawsze wypełniamy najpierw lewą stronę
                    left_idx = i
                    if left_idx < len(chunk_keys):
                        left_key = chunk_keys[left_idx]
                        
                        # --- USUNIĘCIE ZER ---
                        raw_left = left_key.strip()
                        display_left = raw_left.lstrip('0') if raw_left.startswith('0') else raw_left
                        
                        run_l = row.cells[0].paragraphs[0].add_run(display_left)
                        run_l.bold = True
                        
                        val_str_l = f"{decor_summary[left_key]:.1f}".replace('.', ',')
                        run_lm = row.cells[1].paragraphs[0].add_run(val_str_l)
                        run_lm.bold = True
                        
                        # Wstawienie opisu (klucz też ma usunięte zera)
                        row.cells[2].text = self.custom_foils.get(display_left, '')

                    # Jeśli lewa (0-29) jest pełna, program zaczyna uzupełniać prawą (30-59)
                    right_idx = i + MAX_ROWS_PER_PAGE
                    if right_idx < len(chunk_keys):
                        right_key = chunk_keys[right_idx]
                        
                        # --- USUNIĘCIE ZER ---
                        raw_right = right_key.strip()
                        display_right = raw_right.lstrip('0') if raw_right.startswith('0') else raw_right
                        
                        run_r = row.cells[3].paragraphs[0].add_run(display_right)
                        run_r.bold = True
                        
                        val_str_r = f"{decor_summary[right_key]:.1f}".replace('.', ',')
                        run_rm = row.cells[4].paragraphs[0].add_run(val_str_r)
                        run_rm.bold = True
                        
                        # Wstawienie opisu
                        row.cells[5].text = self.custom_foils.get(display_right, '')

                    style_decor_row(row)

        try:
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Błąd zapisu: {e}")
            return False

    # --- POMOCNICZA METODA DO WYPEŁNIANIA SEKCJI DEKORACYJNEJ ---
    def _fill_decor_section(self, doc, data_list, side_title=""):
        if not data_list:
            doc.add_paragraph("Brak zleceń dla tej sekcji.")
            return

        current_base_geometry = None
        table = None
        
        widths = (Cm(6.5), Cm(2.5), Cm(3.0), Cm(5.0))
        
        def style_row(row):
            for i, cell in enumerate(row.cells):
                cell.width = widths[i]
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)

        for idx, item in enumerate(data_list, start=1):
            full_article = str(item['geometry'])
            base_geometry = full_article.split('-')[0]

            if base_geometry != current_base_geometry:
                current_base_geometry = base_geometry
                
                h2 = doc.add_heading('', level=2) 
                
                run_geom = h2.add_run(base_geometry)
                run_geom.font.size = Pt(16)
                run_geom.bold = True
                run_geom.underline = True
                run_geom.font.color.rgb = RGBColor(0, 0, 0)
                
                if side_title:
                    run_side = h2.add_run(f" - {side_title}")
                    run_side.bold = True
                    run_side.underline = False
                    run_side.font.color.rgb = RGBColor(0, 0, 0)
                
                table = doc.add_table(rows=0, cols=4)
                table.style = 'Table Grid'
                table.autofit = False

            if table is not None:
                row = table.add_row()
                row_cells = row.cells
                row_cells[0].text = full_article      
                
                # --- USUNIĘCIE ZER TYLKO DLA FOLII NUMERYCZNYCH ---
                raw_id = str(item['idnrk']).strip()
                display_id = raw_id.lstrip('0') if raw_id.startswith('0') else raw_id
                
                run_idx = row_cells[1].paragraphs[0].add_run(display_id)
                run_idx.bold = True
                
                val_str = f"{item['meters']:.1f}".replace('.', ',')
                run_m = row_cells[2].paragraphs[0].add_run(val_str)
                run_m.bold = True
                
                # Szukamy nazwy folii w słowniku
                row_cells[3].text = self.custom_foils.get(display_id, '')
                
                style_row(row)

    # --- POMOCNICZA METODA DO DODAWANIA NUMERACJI STRON W STOPCE ---
    def _add_page_numbering(self, doc, machine_name, shift_info):
        """Dodaje numerację stron w stopce dokumentu z informacją o maszynie."""
        footer = doc.sections[0].footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # --- budujemy dynamiczne pole numeracji stron (Strona X z Y) ---
        run = p.add_run("Strona ")
        self._append_page_number_field(run, "PAGE")
        run = p.add_run(" z ")
        self._append_page_number_field(run, "NUMPAGES")
        
        # --- dodajemy informację o maszynie i zmianie po numerze strony ---        
        footer_text = "  |  " + f"{machine_name}"
        if shift_info:
            footer_text += f" - {shift_info}"
        p.add_run(footer_text)

    # --- POMOCNICZA METODA DO DODAWANIA DYNAMICZNYCH PÓL NUMERACJI STRON ---
    def _append_page_number_field(self, run, field_name):
        """Pomocnicza metoda do wstawiania pól dynamicznych XML (Strona/Suma stron)."""
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = field_name
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
